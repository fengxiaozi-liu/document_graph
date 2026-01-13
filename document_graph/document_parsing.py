from __future__ import annotations

import logging
import mimetypes
from pathlib import Path


logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {
    ".md",
    ".txt",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
}


class DocumentParseError(RuntimeError):
    pass


class UnsupportedDocumentType(DocumentParseError):
    pass


class MissingParserDependency(DocumentParseError):
    pass


def is_supported_extension(ext: str) -> bool:
    return ext.lower() in SUPPORTED_EXTENSIONS


def _read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def guess_mime_type(path: Path, content_type: str | None = None) -> str | None:
    if content_type:
        return content_type
    mime, _ = mimetypes.guess_type(path.name)
    if mime:
        return mime
    ext = path.suffix.lower()
    if ext == ".pdf":
        return "application/pdf"
    if ext == ".docx":
        return "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    if ext == ".md":
        return "text/markdown"
    if ext in {".html", ".htm"}:
        return "text/html"
    if ext == ".txt":
        return "text/plain"
    return None


def _extract_pdf_text(path: Path) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise MissingParserDependency("missing dependency: pypdf") from exc
    reader = PdfReader(str(path))
    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text:
            parts.append(text)
    return "\n\n".join(parts).strip()


def _extract_docx_text(path: Path) -> str:
    try:
        import docx
    except ImportError as exc:
        raise MissingParserDependency("missing dependency: python-docx") from exc
    document = docx.Document(str(path))
    parts: list[str] = []
    for para in document.paragraphs:
        if para.text:
            parts.append(para.text)
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text]
            if cells:
                parts.append(" | ".join(cells))
    return "\n\n".join(parts).strip()


def read_for_chunking(path: Path) -> tuple[str, str]:
    ext = path.suffix.lower()
    if not is_supported_extension(ext):
        raise UnsupportedDocumentType(f"unsupported file type: {ext or 'unknown'}")

    if ext in {".md", ".txt", ".html", ".htm"}:
        return _read_text(path), ext.lstrip(".")
    if ext == ".pdf":
        return _extract_pdf_text(path), "text"
    if ext == ".docx":
        return _extract_docx_text(path), "text"

    raise UnsupportedDocumentType(f"unsupported file type: {ext or 'unknown'}")


def read_for_preview(path: Path) -> tuple[str, str]:
    text, kind = read_for_chunking(path)
    if kind == "md":
        return text, "markdown"
    return text, "text"
